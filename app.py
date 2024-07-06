from docx import Document
from pathlib import Path
import pandas as pd

from tools import get_time_stamp

from dotenv import load_dotenv
import os
import shutil

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# Load environment variables from .env file
load_dotenv()

BASE_PATH = Path(__file__).resolve().parent


def replace_string_in_docx(docx_filename, old_string, new_string, output_path):
    doc = Document(docx_filename)
    file_name = str(docx_filename).split("/")[-1]
    for p in doc.paragraphs:
        for run in p.runs:
            if old_string in run.text:
                run.text = run.text.replace(old_string, new_string)
    doc.save(f"{output_path}_{file_name}")


def read_names(file_path):
    df = pd.read_excel(str(file_path))
    output = {}
    for row in df.itertuples():
        output[row.name] = row.email
    return output


# read word docs
def read_docx(file_path):
    doc = Document(file_path)
    for p in doc.paragraphs:
        print(p.text)


def send_email(mail_user, app_password, mail_receiver):
    # setup the parameters of the message
    password = "yourapppassword"  # replace with your app password
    msg = MIMEMultipart()
    msg['From'] = "youremail@outlook.com"  # replace with your email
    msg['To'] = "recipientemail@gmail.com"  # replace with recipient's email
    msg['Subject'] = "Subject of the Email"

    # add in the message body
    message = "body of the email"
    msg.attach(MIMEText(message, 'plain'))

    #create server
    server = smtplib.SMTP('smtp.office365.com: 587')

    server.starttls()

    # Login Credentials for sending the mail
    server.login(msg['From'], password)

    # send the message via the server.
    server.sendmail(msg['From'], msg['To'], msg.as_string())

    server.quit()

# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.text import MIMEText

# # setup the parameters of the message
# password = "yourapppassword"  # replace with your app password
# msg = MIMEMultipart('alternative')
# msg['From'] = "youremail@outlook.com"  # replace with your email
# msg['To'] = "recipient1@gmail.com, recipient2@gmail.com, recipient3@gmail.com"  # replace with recipient's emails
# msg['Subject'] = "Subject of the Email"

# # create plain text and HTML version of your message
# text = "Hi!\nHow are you?\nHere is the link you wanted:\nhttps://www.python.org"
# html = """\
# <html>
#   <head></head>
#   <body>
#     <p>Hi!<br>
#        How are you?<br>
#        Here is the <a href="https://www.python.org">link</a> you wanted.
#     </p>
#   </body>
# </html>
# """

# # Record the MIME types of both parts - text/plain and text/html.
# part1 = MIMEText(text, 'plain')
# part2 = MIMEText(html, 'html')

# # Attach parts into message container.
# # According to RFC 2046, the last part of a multipart message, in this case
# # the HTML message, is best and preferred.
# msg.attach(part1)
# msg.attach(part2)

# #create server
# server = smtplib.SMTP('smtp.office365.com: 587')

# server.starttls()

# # Login Credentials for sending the mail
# server.login(msg['From'], password)

# # send the message via the server.
# server.sendmail(msg['From'], msg['To'].split(','), msg.as_string())

# server.quit()


if __name__ == "__main__":

    # time_stamp = get_time_stamp()

    # source_res_dir = f"{BASE_PATH}/result"
    source_origin_dir = f"{BASE_PATH}/source"
    # destination_dir = f"{BASE_PATH}/history/{time_stamp}/result"

    # files = os.listdir(source_res_dir)
    # for file in files:
    #     source = os.path.join(source_res_dir, file)
    #     destination = os.path.join(destination_dir, file)

    #     # create the destination directory if it doesn't exist
    #     os.makedirs(os.path.dirname(destination), exist_ok=True)
    #     shutil.move(source, destination)

    # Get environment variables
    # mail_user = os.getenv('mail_user')
    # mail_pass = os.getenv('mail_pass')
    # mail_receiver = os.getenv('mail_receiver')

    # source_cert = BASE_PATH / 'source' / '2024CPD证书.docx'
    # source_people = BASE_PATH / 'source' / 'names.xlsx'

    # print(BASE_PATH)
    # people = read_names(source_people)
    # for name in people.keys():
    #     print(type(name))
    #     replace_string_in_docx(source_cert, 'LEI    yao', name, f'{BASE_PATH}/res/{name}')

    
    presentation_name_docs = f"{source_origin_dir}/2024全年讲座名单.docx"
    content = read_docx(presentation_name_docs)
    print(content)
